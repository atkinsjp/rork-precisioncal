"""
Build the PrecisionCal AI Agent Training Document in both .docx and .pdf.
Output:
  docs/PrecisionCal-AI-Agent-Training.docx
  docs/PrecisionCal-AI-Agent-Training.pdf
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_LEFT


# ------------------------------------------------------------------ CONTENT --
# Single source of truth: a list of "blocks". Each block is one of:
#   ("h1", "Title")
#   ("h2", "Title")
#   ("h3", "Title")
#   ("p",  "Body paragraph")
#   ("ul", ["item 1", "item 2", ...])
#   ("faq", "Q...?", "A...")
#   ("pb",) — page break

BLOCKS = []
B = BLOCKS.append

# -------------------------------------------------------------------- HEADER
B(("title", "PrecisionCal — AI Agent Training Document"))
B(("subtitle", "Comprehensive product reference for customer-support and pre-sale AI agents"))
B(("p",
   "This document is the canonical training source for the GHL AI agent that represents "
   "PrecisionCal on the website. It explains every feature, where to find it inside the app, "
   "how it works, and how to answer the most common user questions. Hand the entire document "
   "to the agent as knowledge — every section is written so a single passage can stand alone "
   "as an answer."))

# -------------------------------------------------------------------- 1. APP
B(("h1", "1. Product overview"))
B(("p",
   "PrecisionCal is a calm, PhD-led iOS nutrition coach. Instead of raw calorie counting, "
   "it pairs a 5-pass AI meal analyzer with a personal Health Protocol generated during "
   "onboarding, weekly Sunday Calibrations that re-tune the user's plan, a barcode product "
   "scanner that flags additives and allergies, a one-tap hydration tracker, and The Sanctuary "
   "— a moderated community space."))
B(("h3", "Platform & technology"))
B(("ul", [
    "iOS (iPhone). Native SwiftUI app, requires iOS 18 or later.",
    "Local data is stored on-device with SwiftData; no account is required to use the app.",
    "AI features run through PrecisionCal's secure proxy. Photos sent for meal analysis are "
    "used only to produce the analysis result and are not retained for training.",
    "The app works in Light Mode by default to preserve the warm, parchment aesthetic.",
]))
B(("h3", "Brand voice"))
B(("p",
   "Warm, literary, and clinical — described internally as 'a Senior PhD Clinical Nutritionist "
   "speaking quietly across a wooden table.' Avoid hype, exclamation marks, or gamified "
   "language when responding on behalf of PrecisionCal."))

# ----------------------------------------------------------------- 2. NAV
B(("h1", "2. Navigation map"))
B(("p", "The app uses a six-tab bar at the bottom of every screen."))
B(("ul", [
    "Today (hexagon icon) — Dashboard with the Vitality Bloom, Today's Focus, and Sunday Calibration.",
    "Meals (fork & knife icon) — Snap or pick a photo to run the 5-pass meal analyzer; review meal history.",
    "Scan (barcode icon) — Frosted Scanner for packaged-food barcodes; flags allergens and additives.",
    "Sanctuary (leaf icon) — Moderated community feed with Steward review.",
    "Water (drop icon) — One-tap hydration logging with the Fluid Glass visualization.",
    "Profile (person icon) — Daily targets, Sunday Calibration history, and reset.",
]))

# ----------------------------------------------------------------- 3. ONBOARD
B(("h1", "3. Onboarding & calibration"))
B(("p",
   "First launch starts a 9–10 step onboarding wizard. Each step uses a one-thumb dial or "
   "tappable chip set, and the flow is fully reversible with a Back button."))
B(("h3", "Onboarding steps in order"))
B(("ul", [
    "Vision — Welcome screen explaining the PrecisionCal philosophy.",
    "Calibration — Age, weight, and primary goal (Lose / Maintain / Gain).",
    "Hydration — Choose a daily water target in millilitres.",
    "Goals — Pick goal tags (e.g. steady energy, lower sugar, build muscle).",
    "Activity — Sedentary, Light, Moderate, Active, or Very active.",
    "Medical history — High-level toggles (e.g. 'Medical conditions').",
    "Conditions — Only shown if Medical conditions was selected; lists specific conditions.",
    "Allergies — Choose allergens to flag in scanned products and meals.",
    "Medication — List ongoing medications relevant to nutrition.",
    "Synthesis (Protocol) — PrecisionCal generates a personal Health Protocol that becomes the basis for daily focus and Sunday Calibrations.",
]))
B(("h3", "How daily targets are computed"))
B(("p",
   "Calorie target uses the Mifflin–St Jeor equation modified by the activity multiplier "
   "(1.25 sedentary, 1.375 light, 1.45 moderate, 1.55 active, 1.725 very active) and the "
   "primary goal: −400 kcal for Lose, +350 kcal for Gain. Protein target = 1.8 g per kg of "
   "body weight. Carbs target = 45% of calories ÷ 4. Fat target = 30% of calories ÷ 9. "
   "Water target is whatever the user chose in the Hydration step."))

# ----------------------------------------------------------------- 4. DASHBOARD
B(("h1", "4. Today tab — Dashboard"))
B(("p",
   "The Today tab is the daily home screen. From top to bottom it contains: app logo and "
   "greeting, Today's Focus card, optional Sunday Calibration card, the Vitality Bloom, "
   "the optional PhD breakdown drawer, the Hydration card, and recent meals."))

B(("h3", "Today's Focus"))
B(("ul", [
    "A single 6+ word directive composed each day by the AI based on the user's Health Protocol, yesterday's stats, and today's hydration.",
    "Tap the card to instantly cycle to the next curated focus — useful when the user wants a different angle for the day.",
    "If the AI returns a fragment, the app falls back to a curated directive so the card is never empty.",
]))

B(("h3", "Vitality Bloom"))
B(("ul", [
    "A circular bloom that shows three rings: Hydration, Macros, and Adherence.",
    "The center number is calories consumed today versus the calorie target.",
    "Tap the bloom (or swipe up) to open the PhD Breakdown drawer with metabolic impact, nutrient density, and any hazard flags from scanned products.",
    "Swipe down to close the drawer.",
]))

B(("h3", "Sunday Calibration card"))
B(("ul", [
    "Appears Sunday-on-the-week or whenever a new calibration is generated.",
    "Shows three Protocol Pivots — focused adjustments to the user's plan for the upcoming week.",
    "Buttons: 'Show pivots' / 'Collapse' to expand or hide the list, and 'Begin the week' to acknowledge and dismiss.",
    "Acknowledged calibrations remain visible in Profile → View past calibrations.",
]))

B(("h3", "Hydration card"))
B(("p",
   "Compact ring showing today's hydration as a percentage of target. The total includes "
   "explicit water entries plus the water content estimated by the meal analyzer."))

B(("h3", "Recent meals"))
B(("p",
   "Up to four of today's completed meals. Each row shows the title, time, calories, and a "
   "0–100 meal score badge (green ≥ 80, amber 60–79, terracotta < 60)."))

# ----------------------------------------------------------------- 5. MEALS
B(("h1", "5. Meals tab — 5-pass photo analyzer"))
B(("p",
   "Tap 'Choose photo' on the Meals tab, then pick a photo of the meal. PrecisionCal runs "
   "five sequential AI passes and streams progress while the user waits."))
B(("h3", "The five passes"))
B(("ul", [
    "Pass 1 — Identify ingredients and a working title.",
    "Pass 2 — Estimate weight (grams) per item.",
    "Pass 3 — Map to nutrient values (calories, protein, carbs, fat, fibre, sugar, water).",
    "Pass 4 — Synthesize the meal: total scores, metabolic impact summary, QC notes.",
    "Pass 5 — Lipid-sheen scan: visually detects oily sheen and adds a note when present.",
]))
B(("h3", "Reviewing a meal"))
B(("ul", [
    "After analysis, the Meal Analysis sheet opens with full breakdown, per-item nutrients, lipid note, and QC notes.",
    "Tap any past meal in the History list to reopen its sheet.",
    "Long-press a meal to delete it; deletes are immediate and cannot be undone.",
]))
B(("h3", "Failure handling"))
B(("p",
   "If analysis fails (network drop, unreadable image, etc.) the meal is saved with status "
   "'failed' and an error message is shown. The user can delete it and try a clearer photo. "
   "For best results: good lighting, plate roughly fills the frame, top-down or 30° angle."))

# ----------------------------------------------------------------- 6. SCAN
B(("h1", "6. Scan tab — Frosted Scanner"))
B(("p",
   "A fullscreen camera viewfinder with a frosted overlay. Center any EAN, UPC, QR, or "
   "Code128 barcode in the frame. On a successful scan, a soft white 'milk ripple' animation "
   "plays and the Product Detail sheet slides up."))
B(("h3", "What the Product Detail sheet shows"))
B(("ul", [
    "Product name, brand, image, and nutrition facts pulled from the global product database.",
    "Allergy flags cross-referenced against the user's allergies — items the user is allergic to are highlighted in terracotta.",
    "Additive risk: low / medium / high based on the additive panel.",
    "Conflict notes against the user's goals (e.g. 'High sugar conflicts with your sugar-reduction goal').",
]))
B(("h3", "Camera availability"))
B(("p",
   "If the device has no camera (e.g. running in a simulator) the scanner shows a placeholder "
   "asking the user to install the app on their device."))

# ----------------------------------------------------------------- 7. SANCTUARY
B(("h1", "7. Sanctuary tab — Moderated community"))
B(("p",
   "A quiet, PhD-led circle. Every post is read by the Steward (an AI moderator) before it "
   "appears in the public feed. Posts violating community standards are flagged and never go "
   "live."))
B(("h3", "Posting (the Speak FAB)"))
B(("ul", [
    "Tap the terracotta 'Speak' button at the bottom-right to open the composer.",
    "Choose a post kind: Bloom (share today's metrics), Encouragement (share a moment), or Meal Analysis (share a meal score).",
    "Posts begin in a 'reviewing' state, then move to approved or flagged.",
]))
B(("h3", "Reactions and moderation"))
B(("ul", [
    "Heart a post to add encouragement; counts are public.",
    "Comment via the comment sheet; comments are also stewarded.",
    "Report a post via the report icon — reported posts are immediately marked as flagged and removed from the public feed pending review.",
]))
B(("h3", "Founder Dashboard (advanced)"))
B(("p",
   "Power-users can long-press the 'THE SANCTUARY' header for ~1.2 seconds to toggle Founder "
   "mode. While Founder mode is on, a key icon appears in the top-right that opens the "
   "Founder Dashboard with the Review Queue and weekly Innovation Insights aggregated from "
   "community trends. Most consumers will never touch this — it's primarily for the team."))

# ----------------------------------------------------------------- 8. WATER
B(("h1", "8. Water tab — Hydration"))
B(("ul", [
    "A Fluid Glass shows today's water level versus the daily target.",
    "Three one-tap bubbles add 8 oz (237 ml), 12 oz (355 ml), or 16 oz (473 ml).",
    "Each tap gives haptic feedback and animates the glass.",
    "Today's Sips list shows the most recent six entries; swipe left on any entry to delete it.",
]))
B(("p",
   "The Today tab's hydration ring also includes water content from analyzed meals, so users "
   "who eat hydrating foods (soup, fruit) will see total water rise without explicit logging."))

# ----------------------------------------------------------------- 9. PROFILE
B(("h1", "9. Profile tab"))
B(("ul", [
    "Profile card — Primary goal, age, and weight.",
    "Daily Targets — Calories, protein, carbs, fat, water.",
    "Sunday Calibration — Re-run calibration manually (requires at least 3 logged meals in the past 7 days). Shows count of past Protocol Pivots and a link to the full history.",
    "Restart calibration — Wipes the user profile and restarts onboarding. Meals, water entries, and scanned products are preserved.",
]))

# ----------------------------------------------------------------- 10. PRIVACY
B(("h1", "10. Privacy & data handling"))
B(("ul", [
    "All structured data — meals, water, scanned products, profile — is stored on-device via SwiftData.",
    "Meal photos are sent to PrecisionCal's AI proxy only at the moment of analysis. They are processed and discarded; they are not stored on PrecisionCal servers and are not used for AI training.",
    "Barcode lookups send only the barcode number to the global product database.",
    "Sanctuary posts are sent to the Steward (server-side AI moderator) before going public.",
    "There is no account system today; uninstalling the app removes all local data.",
    "The full privacy policy is published at the project website (privacy.html).",
]))

# ----------------------------------------------------------------- 11. SUPPORT
B(("h1", "11. Troubleshooting playbook"))
B(("h3", "Meal analysis fails immediately"))
B(("ul", [
    "Confirm the device has internet connectivity — analysis requires a live AI call.",
    "Try a clearer photo with even lighting and the plate roughly filling the frame.",
    "Delete the failed meal, restart the app, and try again.",
    "If it persists, ask the user to send the failing photo and the error message shown in the alert.",
]))
B(("h3", "Today's Focus shows the same fragment twice"))
B(("ul", [
    "Tap the Focus card to instantly cycle to a new curated directive.",
    "The fragment-rejection guard requires at least 6 words and clean punctuation; rare AI hiccups slip through and are corrected on the next refresh.",
]))
B(("h3", "Scanner doesn't react to a barcode"))
B(("ul", [
    "Ensure the entire barcode (including quiet zones) is inside the viewfinder.",
    "Try better lighting or a steadier hand; reflective packaging can require an angle.",
    "Confirm camera permission is granted for PrecisionCal in iOS Settings.",
]))
B(("h3", "Sunday Calibration button is greyed out"))
B(("ul", [
    "Recalibration requires at least 3 completed meals in the past 7 days.",
    "Log 3 meals via the Meals tab and the button will activate.",
]))
B(("h3", "Want to start over"))
B(("ul", [
    "Profile → Restart calibration. The onboarding wizard runs again and a new Health Protocol is generated.",
    "Meals, scans, and water history are preserved by default. To wipe everything, delete and reinstall the app.",
]))

# ----------------------------------------------------------------- 12. FAQ
B(("h1", "12. Frequently asked questions"))

FAQS = [
    ("What is PrecisionCal?",
     "PrecisionCal is a calm, PhD-led iOS nutrition coach. It analyzes meal photos in five "
     "passes, generates a personal Health Protocol, and recalibrates your plan every Sunday."),

    ("How is PrecisionCal different from a normal calorie counter?",
     "PrecisionCal does not require manual entry. You photograph a meal and the 5-pass "
     "analyzer identifies ingredients, weights, and nutrients automatically. It also assigns "
     "a meal score, detects lipid sheen, and surfaces metabolic impact — context that calorie "
     "counters do not provide."),

    ("Do I need to create an account?",
     "No. PrecisionCal works entirely with on-device data. There is no sign-up, no login, "
     "and no cloud profile."),

    ("Which devices does it support?",
     "iPhone running iOS 18 or later. iPad and Mac are not officially supported today."),

    ("Is there an Android version?",
     "Not yet. Android is on the roadmap but no public date has been announced."),

    ("How accurate is the meal analysis?",
     "Each meal includes QC notes that disclose the analyzer's confidence and any "
     "assumptions it had to make (e.g. estimated portion size). For best results, photograph "
     "meals in good lighting from a top-down or 30° angle."),

    ("Where do my photos go?",
     "They are sent over an encrypted connection to our AI proxy at the moment of analysis, "
     "processed, and discarded. They are not stored on our servers and are not used for AI "
     "training."),

    ("How is my calorie target calculated?",
     "Mifflin–St Jeor equation × your activity multiplier, then ±400 kcal (Lose) or "
     "+350 kcal (Gain) depending on your primary goal. Protein is 1.8 g per kg, carbs are "
     "45% of calories, fat is 30% of calories."),

    ("Can I edit my targets?",
     "Targets recalibrate when you re-run onboarding (Profile → Restart calibration) or when "
     "the Sunday Calibration generates a new Protocol Pivot. Direct manual editing of macro "
     "targets is intentionally not exposed today to keep the experience PhD-led."),

    ("What is the Sunday Calibration?",
     "Once a week (or on demand from the Profile tab) the AI analyzes your last 7 days and "
     "produces three Protocol Pivots — small, focused adjustments to next week's plan."),

    ("What are Protocol Pivots?",
     "Three short, prescriptive recommendations like 'Front-load protein at breakfast' or "
     "'Add a 10-minute post-dinner walk.' Each has a title and a short reason."),

    ("How do I trigger a calibration manually?",
     "Profile tab → 'Re-run calibration.' You need at least 3 completed meals in the last 7 "
     "days for the model to have enough signal."),

    ("What is Today's Focus?",
     "A short daily directive composed from your Health Protocol, yesterday's stats, and "
     "today's hydration. Tap the card to cycle to a new focus instantly."),

    ("What is the Vitality Bloom?",
     "The flower-like ring on the Today tab. The three petals represent Hydration, Macros, "
     "and Adherence. The center number is calories consumed of the daily target."),

    ("What is the PhD Breakdown drawer?",
     "Hidden under the Vitality Bloom. Tap the bloom or swipe up to reveal metabolic impact, "
     "nutrient density score, and any hazard flags from recently scanned products."),

    ("How do I log a meal?",
     "Tap the Meals tab → 'Choose photo' → pick a photo. The 5-pass analyzer runs and the "
     "result opens in a sheet."),

    ("Why did my meal analysis fail?",
     "Usually a network issue or an unreadable image. Confirm internet, retake the photo "
     "with better lighting, delete the failed meal, and try again. Failures are saved so you "
     "can see the error message."),

    ("What is the lipid-sheen scan?",
     "Pass 5 of the meal analyzer. It looks for visible oily sheen on the food and adds a "
     "note when present, since added fats can change the meal's metabolic impact."),

    ("Can I edit a meal's items after analysis?",
     "Manual item-level editing is not exposed today. If the result is wrong you can delete "
     "the meal and re-photograph from a clearer angle."),

    ("Can I log a meal without a photo?",
     "Not currently. The 5-pass analyzer is photo-driven by design. Manual entry is on the "
     "roadmap."),

    ("What does the meal score mean?",
     "A 0–100 score combining nutrient density, macro balance, and lipid context. Green "
     "≥ 80, amber 60–79, terracotta below 60."),

    ("How do I scan a product?",
     "Open the Scan tab and center the barcode in the frame. The Frosted Scanner reads EAN, "
     "UPC, QR, and Code128. On success a milk-ripple animation plays and the Product Detail "
     "sheet appears."),

    ("What does the scanner flag?",
     "Allergens that match your profile, additive risk level, and conflicts with your stated "
     "goals (e.g. high sugar when your goal is sugar reduction)."),

    ("Can I add a product manually if it has no barcode?",
     "Not today. The Frosted Scanner is barcode-driven. Manual product entry is on the "
     "roadmap."),

    ("How do I add water?",
     "Water tab → tap the 8 oz, 12 oz, or 16 oz bubble. Each tap is haptic and animated. "
     "Hydration from analyzed meals is added automatically to your Today tab ring."),

    ("How do I delete a water entry?",
     "Water tab → swipe left on any entry in 'Today's Sips' and tap Delete."),

    ("What is The Sanctuary?",
     "A moderated community feed where members share Bloom updates, Encouragement, and "
     "Meal Analyses. Every post is reviewed by the Steward before it goes public."),

    ("How do I post in the Sanctuary?",
     "Tap the 'Speak' button at the bottom-right of the Sanctuary tab, choose a post kind, "
     "write your message, and submit. The post enters a reviewing state until the Steward "
     "approves it."),

    ("How do I report a post?",
     "Tap the report icon on the post card. Reported posts are immediately flagged and "
     "removed from the public feed pending review."),

    ("What is Founder mode?",
     "An advanced view used by the team. It is hidden from regular users behind a 1.2-second "
     "long-press on the 'THE SANCTUARY' header. End users will not normally encounter it."),

    ("Is my data shared with anyone?",
     "Meals, water, and profile data live on your device. Sanctuary posts are sent to the "
     "Steward for moderation. Barcode lookups send only the barcode to the global product "
     "database. No data is sold."),

    ("How do I delete my data?",
     "Uninstall the app. All on-device data is removed. To restart only the profile, use "
     "Profile → Restart calibration."),

    ("Does PrecisionCal work offline?",
     "Logging water and reviewing existing meals works offline. Meal analysis, barcode "
     "lookup, daily focus, and Sunday Calibration require an internet connection."),

    ("Is PrecisionCal a medical device?",
     "No. PrecisionCal is a wellness and nutrition tool. It is not a substitute for advice "
     "from a licensed clinician. Users with medical conditions should follow their "
     "clinician's plan first."),

    ("Can I use PrecisionCal during pregnancy or while managing a chronic condition?",
     "PrecisionCal can capture context (medical history, conditions, medications) during "
     "onboarding, and the Health Protocol is generated with that context in mind. However, "
     "users with medical conditions should treat the app as supplementary and confirm any "
     "changes with their clinician."),

    ("Is there a subscription or in-app purchase?",
     "Pricing is set on the website and App Store listing. The agent should defer to the "
     "live pricing page rather than quoting a number that may be out of date."),

    ("How do I contact support?",
     "Use the support page on the PrecisionCal website (support.html). Include the iOS "
     "version, the failing photo or barcode if relevant, and a short description of the "
     "issue."),

    ("Where is the privacy policy?",
     "Available on the website at privacy.html. The cookies policy is at cookies.html and "
     "the terms of service at terms.html."),

    ("Will there be an Apple Watch app?",
     "Apple Watch and widget targets are technically supported by the project's architecture "
     "but no public release date has been announced."),

    ("Can I export my data?",
     "Data export is on the roadmap. Today users can screenshot the Profile and Today tabs "
     "for personal records."),
]

for q, a in FAQS:
    B(("faq", q, a))

# -------------------------------------------------------------- 13. CHEAT SHEET
B(("h1", "13. One-paragraph quick-reference for the agent"))
B(("p",
   "PrecisionCal is an iOS (18+) nutrition coach with six tabs: Today (Vitality Bloom + "
   "Today's Focus + Sunday Calibration), Meals (5-pass photo analyzer with lipid-sheen scan), "
   "Scan (Frosted barcode scanner with allergy and additive flags), Sanctuary (Steward-moderated "
   "community), Water (one-tap hydration), and Profile (targets, calibration history, restart). "
   "All structured data is on-device via SwiftData; meal photos go to the AI proxy only at "
   "analysis time and are not retained. Targets are computed from Mifflin–St Jeor and refined "
   "weekly by Sunday Calibration. Onboarding generates a personal Health Protocol that drives "
   "every AI feature. There is no account, no manual macro editing, and no Android version yet."))


# ------------------------------------------------------------ DOCX RENDERER --

ACCENT = RGBColor(0xB5, 0x52, 0x39)         # terracotta
TEXT   = RGBColor(0x2A, 0x21, 0x1B)
MUTED  = RGBColor(0x6B, 0x5E, 0x55)


def _set_run(run, *, size=11, bold=False, italic=False, color=TEXT, font="Helvetica"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def build_docx(path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    for block in BLOCKS:
        kind = block[0]

        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(block[1])
            _set_run(r, size=26, bold=True, color=TEXT)
        elif kind == "subtitle":
            p = doc.add_paragraph()
            r = p.add_run(block[1])
            _set_run(r, size=12, italic=True, color=ACCENT)
            doc.add_paragraph()
        elif kind == "h1":
            p = doc.add_paragraph()
            r = p.add_run(block[1])
            _set_run(r, size=18, bold=True, color=ACCENT)
        elif kind == "h2":
            p = doc.add_paragraph()
            r = p.add_run(block[1])
            _set_run(r, size=14, bold=True, color=TEXT)
        elif kind == "h3":
            p = doc.add_paragraph()
            r = p.add_run(block[1])
            _set_run(r, size=12, bold=True, color=TEXT)
        elif kind == "p":
            p = doc.add_paragraph()
            r = p.add_run(block[1])
            _set_run(r, size=11, color=TEXT)
        elif kind == "ul":
            for item in block[1]:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(item)
                _set_run(r, size=11, color=TEXT)
        elif kind == "faq":
            _, q, a = block
            p = doc.add_paragraph()
            r = p.add_run("Q. " + q)
            _set_run(r, size=11, bold=True, color=ACCENT)
            p = doc.add_paragraph()
            r = p.add_run("A. " + a)
            _set_run(r, size=11, color=TEXT)
            doc.add_paragraph()
        elif kind == "pb":
            doc.add_page_break()

    doc.save(str(path))


# -------------------------------------------------------------- PDF RENDERER --

def build_pdf(path: Path):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title="PrecisionCal AI Agent Training",
        author="PrecisionCal",
    )

    styles = getSampleStyleSheet()
    accent = HexColor("#B55239")
    text   = HexColor("#2A211B")
    muted  = HexColor("#6B5E55")

    s_title = ParagraphStyle("title", parent=styles["Title"],
                             fontName="Helvetica-Bold", fontSize=24,
                             textColor=text, alignment=TA_LEFT, spaceAfter=4)
    s_sub   = ParagraphStyle("sub", parent=styles["Normal"],
                             fontName="Helvetica-Oblique", fontSize=12,
                             textColor=accent, spaceAfter=18)
    s_h1    = ParagraphStyle("h1", parent=styles["Heading1"],
                             fontName="Helvetica-Bold", fontSize=17,
                             textColor=accent, spaceBefore=14, spaceAfter=8,
                             keepWithNext=True)
    s_h2    = ParagraphStyle("h2", parent=styles["Heading2"],
                             fontName="Helvetica-Bold", fontSize=13,
                             textColor=text, spaceBefore=10, spaceAfter=6,
                             keepWithNext=True)
    s_h3    = ParagraphStyle("h3", parent=styles["Heading3"],
                             fontName="Helvetica-Bold", fontSize=11.5,
                             textColor=text, spaceBefore=8, spaceAfter=4,
                             keepWithNext=True)
    s_body  = ParagraphStyle("body", parent=styles["BodyText"],
                             fontName="Helvetica", fontSize=10.5,
                             leading=15, textColor=text, spaceAfter=8)
    s_bullet = ParagraphStyle("bullet", parent=s_body,
                              fontName="Helvetica", fontSize=10.5,
                              leading=14, leftIndent=14, spaceAfter=2)
    s_q     = ParagraphStyle("q", parent=s_body,
                             fontName="Helvetica-Bold", fontSize=10.5,
                             textColor=accent, spaceAfter=2,
                             keepWithNext=True)
    s_a     = ParagraphStyle("a", parent=s_body,
                             fontName="Helvetica", fontSize=10.5,
                             leading=14, textColor=text, spaceAfter=10)

    def esc(t: str) -> str:
        return (t.replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;"))

    flow = []
    for block in BLOCKS:
        kind = block[0]
        if kind == "title":
            flow.append(Paragraph(esc(block[1]), s_title))
        elif kind == "subtitle":
            flow.append(Paragraph(esc(block[1]), s_sub))
        elif kind == "h1":
            flow.append(Paragraph(esc(block[1]), s_h1))
        elif kind == "h2":
            flow.append(Paragraph(esc(block[1]), s_h2))
        elif kind == "h3":
            flow.append(Paragraph(esc(block[1]), s_h3))
        elif kind == "p":
            flow.append(Paragraph(esc(block[1]), s_body))
        elif kind == "ul":
            items = [ListItem(Paragraph(esc(i), s_bullet), leftIndent=10,
                              bulletColor=accent) for i in block[1]]
            flow.append(ListFlowable(
                items, bulletType="bullet", start="•",
                bulletFontName="Helvetica-Bold", bulletFontSize=10,
                leftIndent=14, bulletColor=accent
            ))
            flow.append(Spacer(1, 6))
        elif kind == "faq":
            _, q, a = block
            flow.append(Paragraph("Q. " + esc(q), s_q))
            flow.append(Paragraph("A. " + esc(a), s_a))
        elif kind == "pb":
            flow.append(PageBreak())

    doc.build(flow)


# ----------------------------------------------------------------- MAIN --
def main():
    out_dir = Path("docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "PrecisionCal-AI-Agent-Training.docx"
    pdf_path  = out_dir / "PrecisionCal-AI-Agent-Training.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(f"Wrote {docx_path} ({docx_path.stat().st_size} bytes)")
    print(f"Wrote {pdf_path} ({pdf_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
